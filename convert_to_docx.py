#!/usr/bin/env python3
"""
Simple Markdown to DOCX converter for user_stories.md
This script converts the markdown file to a Word document with preserved formatting.
"""

import re
import os

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Error: python-docx is not installed.")
    print("Please run: pip3 install python-docx --user")
    exit(1)

def convert_md_to_docx(md_file, docx_file):
    """Convert markdown file to DOCX with formatting"""
    
    # Create a new Document
    doc = Document()
    
    # Read the markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Main title (# )
        if line.startswith('# '):
            title = line[2:].strip()
            p = doc.add_heading(title, level=1)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
        # Section headers (## )
        elif line.startswith('## '):
            header = line[3:].strip()
            doc.add_heading(header, level=2)
            
        # Subsection headers (### )
        elif line.startswith('### '):
            subheader = line[4:].strip()
            doc.add_heading(header, level=3)
            
        # Horizontal rule (---)
        elif line.strip() == '---':
            doc.add_paragraph('_' * 80)
            
        # Bold text (**text**)
        elif line.startswith('**') and line.endswith('**'):
            text = line[2:-2]
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
            
        # Numbered list
        elif re.match(r'^\d+\.', line):
            doc.add_paragraph(line, style='List Number')
            
        # Bullet list with dash
        elif line.startswith('- '):
            # Handle bold within list items
            text = line[2:]
            p = doc.add_paragraph(style='List Bullet')
            
            # Parse inline bold
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
                    
        # Indented list (starts with spaces)
        elif line.startswith('  - ') or line.startswith('   -'):
            text = line.strip()[2:]
            p = doc.add_paragraph(text, style='List Bullet 2')
            
        # Regular paragraph
        else:
            # Parse inline bold
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*|`.*?`)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.color.rgb = RGBColor(200, 0, 0)
                else:
                    p.add_run(part)
        
        i += 1
    
    # Save the document
    doc.save(docx_file)
    print(f"✅ Successfully converted {md_file} to {docx_file}")
    print(f"📄 Document location: {os.path.abspath(docx_file)}")

if __name__ == "__main__":
    md_file = "user_stories.md"
    docx_file = "user_stories.docx"
    
    if not os.path.exists(md_file):
        print(f"Error: {md_file} not found!")
        exit(1)
    
    convert_md_to_docx(md_file, docx_file)
